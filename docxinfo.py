# import flask
import docx 
import json,os,re
# 导入pymysql模块
import pymysql
def getText(path):
    doc = docx.Document(path) #读入文件
    dict={}
    count=1
    dict['q'+str(count)]={}
    dict['q'+str(count)]["question"]=""
    for i in range(len(doc.paragraphs)):
        print(i,doc.paragraphs[i].text)
        if doc.paragraphs[i].text == "":
            continue
        elif  doc.paragraphs[i].text[0]=='A':
            dict['q'+str(count)]['opA']=doc.paragraphs[i].text

        elif  doc.paragraphs[i].text[0]=='B':
            dict['q'+str(count)]['opB']=doc.paragraphs[i].text

        elif  doc.paragraphs[i].text[0]=='C':
            dict['q'+str(count)]['opC']=doc.paragraphs[i].text
        
        elif  doc.paragraphs[i].text[0]=='D':
            dict['q'+str(count)]['opD']=doc.paragraphs[i].text
            count+=1
            dict['q'+str(count)]={}
            dict['q'+str(count)]["question"]=""
            print(dict)
        elif  doc.paragraphs[i].text[0]=='E':
            print('q'+str(count-1),count,count-1)
            dict['q'+str(count-1)]['opE']=doc.paragraphs[i].text
        else:
            dict['q'+str(count)]["question"]+=doc.paragraphs[i].text
    return dict

# dict_rel = doc.part._rels
# for rel in dict_rel:
#     rel = dict_rel[rel]
#     print(rel,rel.target_ref)
#     if "image" in rel.target_ref:
#         img_name = re.findall("/(.*)", rel.target_ref)[0]
#         result_path='img/'
#         with open(f'{result_path}{img_name}', "wb") as f:
#             f.write(rel.target_part.blob)

def tojson(file,data):
    with open(file,'w+',encoding='utf-8') as f:
        json.dump(data,f,ensure_ascii=False,indent=4) 
    print('生成json数据成功，文件位置：',file)

def jsontomysql(file):
    with open(file,'r',encodinginfo='utf-8') as f:
        data=json.load(f)
    # conn = pymysql.connect(
    #     host='localhost', port=3306,user='root',password='ccb12345',database='test',charset='utf8')
    # cursor = conn.cursor()
    with open('data.sql','a+',encoding='utf-8') as f1:
        for item in data.values():
            sql = f"insert into  test.questions(question,opA,opB,opC,opD,opE)values('{item['question']}','{item['opA']}','{item['opB']}','{item['opC']}','{item['opD']}','{item.get('opE','')}');\n"
            print(item['question'])
            f1.write(sql)
        
    #     try: 
    #         # 执行SQL语句
    #         cursor.execute(sql)
    #         # 提交事务
    #         conn.commit()
    #     except Exception as e:
    #         print('dberror')
    #         conn.rollback()
    # cursor.close()
    # # 关闭数据库连接
    # conn.close()
   
        
if __name__ == '__main__':
    # path=r'D:\wangjz\软件资源\计算机四级真题1600.docx'
    # data=getText(path)
    # file=path.split('.')[0]+'.json'
    # print(file)
    # tojson(file,data)
    file='D:\wangjz\软件资源\计算机四级真题1600.json'
    jsontomysql(file)